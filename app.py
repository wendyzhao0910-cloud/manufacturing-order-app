import os
import re
import tempfile
from datetime import datetime

import streamlit as st
from docx import Document

# 可选 PDF
PDF_AVAILABLE = True
try:
    from docx2pdf import convert
except Exception:
    PDF_AVAILABLE = False


# =========================
# 基础工具
# =========================
def today_str_dash():
    """当天日期：2026-3-3"""
    now = datetime.today()
    return f"{now.year}-{now.month}-{now.day}"


def ship_date_to_filename_date(ship_date: str) -> str:
    """出货日期转文件名日期：2026-3-3 -> 2026.3.3"""
    parts = ship_date.strip().split("-")
    if len(parts) == 3:
        return f"{int(parts[0])}.{int(parts[1])}.{int(parts[2])}"
    return ship_date.replace("-", ".")


def sanitize_filename(name: str) -> str:
    return re.sub(r'[\\/:*?"<>|]', "_", str(name).strip())


def normalize_text(text: str) -> str:
    return str(text).replace("\u3000", " ").strip()


def replace_paragraph_text(paragraph, new_text: str):
    """
    直接重写整段，兼容 run 被拆分。
    """
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def clear_paragraph(paragraph):
    replace_paragraph_text(paragraph, "")


def get_table_paragraphs(table):
    paras = []
    for row in table.rows:
        for cell in row.cells:
            paras.extend(cell.paragraphs)
            for t in cell.tables:
                paras.extend(get_table_paragraphs(t))
    return paras


def get_all_body_paragraphs(doc: Document):
    paras = list(doc.paragraphs)
    for table in doc.tables:
        paras.extend(get_table_paragraphs(table))
    return paras


def get_all_header_paragraphs(doc: Document):
    paras = []
    for section in doc.sections:
        paras.extend(section.header.paragraphs)
        for table in section.header.tables:
            paras.extend(get_table_paragraphs(table))
    return paras


# =========================
# 规则1：PI/订单编号：
# 只识别“：”后面的内容替换
# =========================
def replace_pi_no_after_colon(paragraph, new_pi: str):
    text = normalize_text(paragraph.text)
    if not text:
        return False, None

    pattern = r"^(PI/订单编号[：:])(.*)$"
    m = re.match(pattern, text)
    if m:
        old_value = m.group(2).strip()
        new_text = f"{m.group(1)}{new_pi}"
        replace_paragraph_text(paragraph, new_text)
        return True, f"{old_value} -> {new_pi}"

    return False, None


# =========================
# 规则2：出货日期：
# 只识别“：”后面的内容替换
# =========================
def replace_ship_date_after_colon(paragraph, new_date: str):
    text = normalize_text(paragraph.text)
    if not text:
        return False, None

    pattern = r"^(出货日期[：:])(.*)$"
    m = re.match(pattern, text)
    if m:
        old_value = m.group(2).strip()
        new_text = f"{m.group(1)}{new_date}"
        replace_paragraph_text(paragraph, new_text)
        return True, f"{old_value} -> {new_date}"

    return False, None


# =========================
# 规则3：A/B/C/D款数量行
# 只改 “A款黑色：100台”
# 不改 “A款黑色：黑色斜纹阳离子5051+...”
# =========================
def replace_model_qty_line(paragraph, model_letter: str, model_desc: str, qty: int):
    text = normalize_text(paragraph.text)
    if not text:
        return False, None

    pattern = rf"^\s*{re.escape(model_letter)}款.*?[：:]\s*\d+\s*台\s*$"
    if re.match(pattern, text):
        old_text = text
        new_text = f"{model_letter}款{model_desc}：{qty}台"
        replace_paragraph_text(paragraph, new_text)
        return True, f"{old_text} -> {new_text}"

    return False, None


def is_model_qty_line(text: str, model_letter: str):
    text = normalize_text(text)
    pattern = rf"^\s*{re.escape(model_letter)}款.*?[：:]\s*\d+\s*台\s*$"
    return bool(re.match(pattern, text))


def is_model_desc_line(text: str, model_letter: str):
    """
    识别类似：
    B款黑色：黑色斜纹阳离子5051+黑色满天星三明治网布
    即：B款开头、冒号后不是“数字台”
    """
    text = normalize_text(text)
    if not text:
        return False

    if not re.match(rf"^\s*{re.escape(model_letter)}款", text):
        return False

    if "：" not in text and ":" not in text:
        return False

    # 如果是数量行，不算说明行
    if is_model_qty_line(text, model_letter):
        return False

    return True


def delete_unused_model_lines(paragraphs, used_model_letters, logs):
    all_letters = ["A", "B", "C", "D"]
    unused_letters = [x for x in all_letters if x not in used_model_letters]

    for p in paragraphs:
        text = normalize_text(p.text)
        if not text:
            continue

        for letter in unused_letters:
            if is_model_qty_line(text, letter) or is_model_desc_line(text, letter):
                logs.append(f"[删除未使用款] 删除：{text}")
                clear_paragraph(p)
                break


def replace_total_line(paragraph, total_qty: int, total_boxes: int):
    """
    合计 300台/20箱（N）
    只改 台 和 箱 前面的数字
    """
    text = normalize_text(paragraph.text)
    if not text:
        return False, None

    pattern = r"(合计\s*)(\d+)(\s*台\s*/\s*)(\d+)(\s*箱.*)"
    m = re.search(pattern, text)
    if m:
        old_text = text
        new_text = f"{m.group(1)}{total_qty}{m.group(3)}{total_boxes}{m.group(5)}"
        replace_paragraph_text(paragraph, new_text)
        return True, f"{old_text} -> {new_text}"

    return False, None


# =========================
# 规则4：文件中其他含有 KXXXLXXXX 的，一并修改
# 只替换标准格式 K123L4567
# =========================
def replace_all_pi_codes(paragraph, new_pi: str):
    text = normalize_text(paragraph.text)
    if not text:
        return False, None

    old_text = text
    new_text = re.sub(r"\bK\d{3}L\d{4}\b", new_pi, text)

    if new_text != old_text:
        replace_paragraph_text(paragraph, new_text)
        return True, f"{old_text} -> {new_text}"

    return False, None


# =========================
# 规则5：唯一性文件
# 只要有“唯一性”或“序列号”相关就可替换
# 不同种类分别编号
# 不需要“附加说明”输入
# =========================
def is_unique_file_candidate(text: str):
    text = normalize_text(text)
    if not text:
        return False

    keywords = ["唯一性", "序列号"]
    type_words = ["不干胶", "热转印", "贴纸", "标签"]

    if any(k in text for k in keywords):
        return True
    if "序列号" in text and any(t in text for t in type_words):
        return True
    return False


def build_unique_line_by_old_text(old_text: str, file_no: int, file_type: str, new_pi: str, serial_range: str):
    """
    尽量保留原句结构，只替换：
    1. 前面的序号（如28）
    2. 文件种类（不干胶/热转印）
    3. PI相关编码
    4. 序列号范围
    """
    text = old_text

    # 1) 替换开头数字
    text = re.sub(r"^\s*\d+", str(file_no), text, count=1)

    # 2) 如果原文含文件类型词，替换成当前类型；如果没有，尽量在开头数字后插入
    if any(t in text for t in ["不干胶", "热转印", "贴纸", "标签"]):
        text = re.sub(r"(不干胶|热转印|贴纸|标签)", file_type, text, count=1)
    else:
        text = re.sub(r"^(\s*\d+)", rf"\1{file_type}", text, count=1)

    # 3) 替换所有标准 PI 编号
    text = re.sub(r"\bK\d{3}L\d{4}\b", new_pi, text)

    # 4) 对 K265 C004 这种也尽量替换成新 PI
    text = re.sub(r"\bK\d+\s+[A-Z]\d+\b", new_pi, text)

    # 5) 替换序列号范围
    if re.search(r"序列号\s*\d+\s*-\s*\d+", text):
        text = re.sub(r"(序列号\s*)\d+\s*-\s*\d+", rf"\1{serial_range}", text, count=1)
    else:
        # 若没有明显“序列号xxxx-xxxx”，就在“序列号”后补
        text = re.sub(r"(序列号)", rf"\1{serial_range}", text, count=1)

    return text


def replace_unique_file_lines(paragraphs, unique_items, logs):
    """
    unique_items:
    [
      {"file_type":"不干胶","file_no":29,"serial_range":"004661-005100"},
      ...
    ]
    按类型分别匹配和替换。
    """
    # 先收集候选
    candidates = []
    for p in paragraphs:
        t = normalize_text(p.text)
        if is_unique_file_candidate(t):
            candidates.append((p, t))

    # 按种类分候选
    typed_candidates = {}
    for p, t in candidates:
        matched = False
        for tp in ["不干胶", "热转印"]:
            if tp in t:
                typed_candidates.setdefault(tp, []).append((p, t))
                matched = True
                break
        if not matched:
            typed_candidates.setdefault("其他唯一性文件", []).append((p, t))

    # 按输入项目类型分组
    items_by_type = {}
    for item in unique_items:
        items_by_type.setdefault(item["file_type"], []).append(item)

    for file_type, items in items_by_type.items():
        targets = typed_candidates.get(file_type, [])
        for i, item in enumerate(items):
            if i < len(targets):
                p, old_text = targets[i]
                new_text = build_unique_line_by_old_text(
                    old_text=old_text,
                    file_no=item["file_no"],
                    file_type=item["file_type"],
                    new_pi=item["new_pi"],
                    serial_range=item["serial_range"],
                )
                replace_paragraph_text(p, new_text)
                logs.append(f"[唯一性文件-{file_type}] {old_text} -> {new_text}")
            else:
                logs.append(f"[唯一性文件-{file_type}] 模板中未找到第 {i+1} 条可替换位置，请检查模板。")


# =========================
# 规则6：页眉日期改为当天
# =========================
def replace_header_date(paragraph, new_date: str):
    text = normalize_text(paragraph.text)
    if not text:
        return False, None

    pattern = r"\b\d{4}[-/\.]\d{1,2}[-/\.]\d{1,2}\b"
    m = re.search(pattern, text)
    if m:
        old_value = m.group(0)
        new_text = re.sub(pattern, new_date, text, count=1)
        replace_paragraph_text(paragraph, new_text)
        return True, f"{old_value} -> {new_date}"

    return False, None


# =========================
# 规则7：输出文件名
# 只改：
# 1. 订单编号 KXXXLXXXX
# 2. 后面的日期，改成出货日期格式 2026.4.30
# 文件名示例：
# K265L5313 日本LEAMAN C004 黑色 大货制造令2026.4.30 R0
# =========================
def generate_output_filename(uploaded_filename: str, new_pi: str, ship_date: str):
    base = os.path.splitext(uploaded_filename)[0]
    ship_date_filename = ship_date_to_filename_date(ship_date)

    # 替换订单编号
    new_base = re.sub(r"\bK\d{3}L\d{4}\b", new_pi, base)

    # 替换文件名中的日期（2026.4.30 / 2026-4-30 / 2026/4/30）
    new_base2 = re.sub(
        r"\b\d{4}[-/\.]\d{1,2}[-/\.]\d{1,2}\b",
        ship_date_filename,
        new_base
    )

    # 如果原文件名里没有订单编号，则前置
    if new_base2 == base:
        new_base2 = f"{new_pi} {base}"

    return sanitize_filename(new_base2) + ".docx"


# =========================
# PDF
# =========================
def convert_to_pdf(docx_path: str):
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    convert(docx_path, pdf_path)
    return pdf_path


# =========================
# 主处理逻辑
# =========================
def process_document(uploaded_file, data):
    logs = []

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(uploaded_file.read())
        temp_path = tmp.name

    doc = Document(temp_path)

    body_paragraphs = get_all_body_paragraphs(doc)
    header_paragraphs = get_all_header_paragraphs(doc)

    # 1. PI/订单编号：
    for p in body_paragraphs:
        changed, detail = replace_pi_no_after_colon(p, data["pi_no"])
        if changed:
            logs.append(f"[PI/订单编号] {detail}")

    # 2. 出货日期：
    for p in body_paragraphs:
        changed, detail = replace_ship_date_after_colon(p, data["ship_date"])
        if changed:
            logs.append(f"[出货日期] {detail}")

    # 3. 款式数量
    total_qty = 0
    used_model_letters = []

    for item in data["models"]:
        if item["enabled"] and item["model_desc"].strip() and item["qty"] > 0:
            total_qty += item["qty"]
            used_model_letters.append(item["model_letter"])

            for p in body_paragraphs:
                changed, detail = replace_model_qty_line(
                    p,
                    model_letter=item["model_letter"],
                    model_desc=item["model_desc"],
                    qty=item["qty"]
                )
                if changed:
                    logs.append(f"[{item['model_letter']}款数量] {detail}")

    # 删除未使用款式的数量行和说明行
    delete_unused_model_lines(body_paragraphs, used_model_letters, logs)

    # 合计
    for p in body_paragraphs:
        changed, detail = replace_total_line(p, total_qty, data["total_boxes"])
        if changed:
            logs.append(f"[合计] {detail}")

    # 4. 其他所有 KXXXLXXXX 一并修改
    for p in body_paragraphs:
        changed, detail = replace_all_pi_codes(p, data["pi_no"])
        if changed:
            logs.append(f"[全文订单编号替换] {detail}")

    # 5. 唯一性文件
    if data["has_unique_files"] and data["unique_files"]:
        replace_unique_file_lines(body_paragraphs, data["unique_files"], logs)

    # 6. 页眉日期
    for p in header_paragraphs:
        changed, detail = replace_header_date(p, data["today_date"])
        if changed:
            logs.append(f"[页眉日期] {detail}")

    # 7. 输出文件名自动改订单编号和日期
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)

    output_docx_name = generate_output_filename(
        uploaded_filename=uploaded_file.name,
        new_pi=data["pi_no"],
        ship_date=data["ship_date"]
    )
    output_docx_path = os.path.join(output_dir, output_docx_name)

    doc.save(output_docx_path)

    if os.path.exists(temp_path):
        os.remove(temp_path)

    return output_docx_path, logs


# =========================
# Streamlit 页面
# =========================
st.set_page_config(page_title="翻单制造令生成工具", page_icon="📄")
st.title("翻单制造令生成工具")

uploaded_template = st.file_uploader("上传制造令模板（.docx）", type=["docx"])

today_date_value = today_str_dash()

st.subheader("基础信息")
col1, col2 = st.columns(2)

with col1:
    pi_no = st.text_input("1. PI/订单编号（格式：KXXXLXXXX）", "")
    ship_date = st.text_input("2. 出货日期（格式：2026-3-3）", "")

with col2:
    st.text_input("6. 页眉日期（自动为当天）", value=today_date_value, disabled=True)
    total_boxes = st.number_input("合计箱数", min_value=0, step=1, value=0)

if pi_no and not re.fullmatch(r"K\d{3}L\d{4}", pi_no):
    st.warning("PI/订单编号建议格式：KXXXLXXXX，例如 K265L5313")

if ship_date and not re.fullmatch(r"\d{4}-\d{1,2}-\d{1,2}", ship_date):
    st.warning("出货日期建议格式：2026-3-3")

st.subheader("3. 款式及数量")
model_letters = ["A", "B", "C", "D"]
models = []

for letter in model_letters:
    st.markdown(f"**{letter}款**")
    c1, c2, c3 = st.columns([1, 3, 1])

    with c1:
        enabled = st.checkbox(f"启用{letter}款", value=(letter == "A"), key=f"enabled_{letter}")

    with c2:
        model_desc = st.text_input(
            f"{letter}款颜色/描述",
            "",
            key=f"desc_{letter}",
            placeholder="例如：黑色 / 红色 / 黑色+灰边"
        )

    with c3:
        qty = st.number_input(
            f"{letter}款台数",
            min_value=0,
            step=1,
            key=f"qty_{letter}"
        )

    models.append({
        "model_letter": letter,
        "enabled": enabled,
        "model_desc": model_desc,
        "qty": int(qty),
    })

st.subheader("5. 唯一性文件")
has_unique_files = st.checkbox("是否有唯一性文件", value=False)
unique_files = []

if has_unique_files:
    st.write("不同种类的唯一性文件，编号分别独立统计。")

    unique_type_count = st.number_input("唯一性文件种类数量", min_value=1, step=1, value=1)

    for t in range(int(unique_type_count)):
        st.markdown(f"### 唯一性文件种类 {t+1}")
        c1, c2, c3 = st.columns(3)

        with c1:
            file_type = st.selectbox(
                f"文件种类 {t+1}",
                ["不干胶", "热转印", "其他唯一性文件"],
                key=f"file_type_{t}"
            )

        with c2:
            item_count = st.number_input(
                f"{file_type}数量 {t+1}",
                min_value=1,
                step=1,
                value=1,
                key=f"item_count_{t}"
            )

        with c3:
            start_no = st.number_input(
                f"{file_type}起始编号 {t+1}",
                min_value=1,
                step=1,
                value=29,
                key=f"start_no_{t}"
            )

        for i in range(int(item_count)):
            st.markdown(f"**{file_type} 第 {i+1} 条**")
            c4, c5 = st.columns(2)

            with c4:
                serial_range = st.text_input(
                    f"{file_type} 序列号范围 {i+1}",
                    key=f"serial_{t}_{i}",
                    placeholder="例如：004661-005100"
                )

            with c5:
                st.text_input(
                    f"{file_type} 文件编号 {i+1}",
                    value=pi_no,
                    key=f"file_code_show_{t}_{i}",
                    disabled=True
                )

            unique_files.append({
                "file_type": file_type,
                "file_no": int(start_no) + i,
                "serial_range": serial_range.strip(),
                "new_pi": pi_no.strip(),
            })

generate_pdf = st.checkbox("同时生成 PDF", value=False)

if st.button("生成翻单制造令"):
    if uploaded_template is None:
        st.error("请先上传 Word 模板。")
    elif not pi_no.strip():
        st.error("请填写 PI/订单编号。")
    elif not ship_date.strip():
        st.error("请填写出货日期。")
    else:
        data = {
            "pi_no": pi_no.strip(),
            "ship_date": ship_date.strip(),
            "today_date": today_date_value,
            "models": models,
            "total_boxes": int(total_boxes),
            "has_unique_files": has_unique_files,
            "unique_files": unique_files,
        }

        try:
            output_docx, logs = process_document(uploaded_template, data)

            st.success(f"Word 已生成：{output_docx}")

            with open(output_docx, "rb") as f:
                st.download_button(
                    label="下载 Word 文件",
                    data=f,
                    file_name=os.path.basename(output_docx),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            if generate_pdf:
                if not PDF_AVAILABLE:
                    st.warning("当前环境未安装 docx2pdf，无法生成 PDF。请先安装：python -m pip install docx2pdf")
                else:
                    try:
                        output_pdf = convert_to_pdf(output_docx)
                        st.success(f"PDF 已生成：{output_pdf}")
                        with open(output_pdf, "rb") as f:
                            st.download_button(
                                label="下载 PDF 文件",
                                data=f,
                                file_name=os.path.basename(output_pdf),
                                mime="application/pdf"
                            )
                    except Exception as e:
                        st.error(f"PDF 生成失败：{e}")

            st.subheader("处理日志")
            if logs:
                for log in logs:
                    st.write("- " + log)
            else:
                st.warning("没有识别到可替换内容，说明模板写法和规则不完全一致，需要按真实模板继续微调。")

        except Exception as e:
            st.error(f"生成失败：{e}")